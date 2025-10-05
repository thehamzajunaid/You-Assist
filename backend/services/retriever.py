from PyPDF2 import PdfReader
from docx import Document
import io

# def extract_text_from_pdf(file_content: bytes) -> str:
#     """Extract text from PDF file"""
#     pdf_reader = PdfReader(io.BytesIO(file_content))
#     text = ""
#     for page in pdf_reader.pages:
#         text += page.extract_text() + "\n"
#     return text

# def extract_text_from_docx(file_content: bytes) -> str:
#     """Extract text from DOCX file"""
#     doc = Document(io.BytesIO(file_content))
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text

# def extract_text_from_txt(file_content: bytes) -> str:
#     """Extract text from TXT file"""
#     return file_content.decode('utf-8')


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error extracting PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return file_content.decode('latin-1')
        except Exception as e:
            raise Exception(f"Error decoding text file: {str(e)}")